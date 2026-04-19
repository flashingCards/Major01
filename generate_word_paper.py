from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "final_research_paper_divash_team_word_ready.docx"


def set_page_margins(section, margin_inch=1.0):
    margin = Inches(margin_inch)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


def set_default_font(document):
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12 if style_name == "Normal" else 14)
        element = style.element
        rpr = element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_centered(document, text, size=12, bold=False, space_after=6):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    return p


def add_paragraph(document, text, italic=False, bold=False):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = italic
    run.bold = bold
    return p


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_number(document, text):
    p = document.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_table(document, title, headers, rows):
    add_paragraph(document, title, bold=True)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    document.add_paragraph()


def add_picture(document, image_path, caption, width=6.2):
    document.add_picture(str(image_path), width=Inches(width))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def build_document():
    doc = Document()
    set_default_font(doc)
    section = doc.sections[0]
    set_page_margins(section)

    add_centered(doc, "Simulation-Based Implementation of Lean Manufacturing Techniques for Productivity Improvement:", size=16, bold=True, space_after=2)
    add_centered(doc, "A Data-Driven Case Study From a Beverage Bottling Line", size=16, bold=True, space_after=18)
    add_centered(doc, "Final Research Paper", size=14, bold=True, space_after=18)
    add_centered(doc, "Submitted by: Divash Krishnam", size=12, bold=True)
    add_centered(doc, "Team Members: Ashutosh Verma, Raja", size=12)
    add_centered(doc, "Department of Production and Industrial Engineering", size=12)
    add_centered(doc, "Guided by: Dr. MOHD SHUAIB", size=12, bold=True)
    add_centered(doc, "Mechanical Department", size=12)
    add_centered(doc, "Date: March 26, 2026", size=12)
    doc.add_page_break()

    add_heading(doc, "Abstract", 1)
    add_paragraph(
        doc,
        "Lean manufacturing is widely recognized as a practical and effective strategy for improving productivity, reducing waste, and stabilizing production flow. However, many small and medium-sized manufacturing firms still face difficulty in deciding which lean techniques should be implemented first, especially when resources are limited and repeated experimentation on the shop floor is risky. This study develops a self-contained, data-driven research framework for evaluating lean interventions using a real industrial production dataset from a beverage bottling line monitored through an Industrial Internet of Things (IIoT) system. The dataset contains 59 production summary records across 56 calendar production days between July 22, 2022 and February 6, 2023, 265 hourly operating observations, and 1,388 downtime events. The study first analyzes the production system through throughput, operating time, pause time, product mix, and temporal efficiency variation. Next, the downtime structure is examined by classifying events into micro-, minor-, and major-stoppage categories. Finally, an event-level bootstrap simulation is used to compare three lean scenarios: 5S with visual management, total productive maintenance (TPM), and an integrated lean bundle. The observed system produced 212,358 liters with an overall availability of 42.66%. Downtime analysis shows that micro-stoppages represent 61.31% of all events but only 13.59% of total downtime minutes, whereas major stoppages represent only 6.77% of events but consume 66.26% of lost time. Simulation results show that 5S and visual management improve output by 2.33%, TPM improves output by 20.61%, and the integrated lean bundle improves output by 32.19%, while increasing simulated availability from 42.37% to 55.33%. The main implication is that maintenance-led reduction of major stoppages should be prioritized before broader workplace-organization efforts when downtime severity is highly concentrated. The study contributes a reproducible open-data framework for simulation-based lean analysis and provides a complete final-year research-paper template for manufacturing engineering students."
    )

    add_heading(doc, "Keywords", 1)
    add_paragraph(doc, "Lean manufacturing, productivity improvement, small-scale industries, discrete-event simulation, beverage bottling line, downtime analysis, total productive maintenance, 5S, IIoT, digital manufacturing.")

    add_heading(doc, "1. Introduction", 1)
    intro_paras = [
        "Lean manufacturing has become one of the most influential improvement philosophies in industrial engineering because it focuses on eliminating waste, improving flow, increasing responsiveness, and enhancing process discipline. In practice, however, lean implementation is rarely straightforward. Many organizations, especially smaller firms, understand the value of lean at a conceptual level but struggle to determine which tools should be applied first and where the highest-impact waste actually exists.",
        "For small-scale industries and SMEs, this decision is especially important. Such firms often operate with constrained budgets, limited technical manpower, and less tolerance for failed improvement initiatives. If a company selects the wrong improvement priority, it may spend time and effort without producing meaningful productivity gains. For this reason, simulation-based lean implementation is highly valuable. It allows researchers and managers to test alternative interventions in a virtual decision environment before changing the real production system.",
        "The topic originally considered for this work was Simulation-Based Implementation of Lean Manufacturing Techniques for Productivity Improvement in Small-Scale Industries. That topic remains highly relevant, but the present paper refines it into a stronger research design by using a real open industrial dataset instead of hypothetical or purely conceptual examples. The use of real production and downtime data makes the analysis more rigorous, reproducible, and suitable for final-year academic work.",
        "The present study uses an open dataset from a beverage bottling line that was monitored through IIoT infrastructure. This is particularly useful because the dataset includes production output, operating time, and detailed downtime-event logs, allowing a direct connection between shop-floor losses and lean scenario design. Instead of merely describing lean concepts, the paper quantitatively tests which lean interventions would provide the highest productivity improvement.",
    ]
    for para in intro_paras:
        add_paragraph(doc, para)
    add_paragraph(doc, "The main objectives of this study are:", bold=True)
    for item in [
        "to analyze the production behavior and downtime structure of a real bottling-line system;",
        "to identify the dominant operational waste pattern affecting productivity;",
        "to simulate the productivity impact of selected lean interventions; and",
        "to recommend an implementation priority for lean techniques based on data-driven evidence.",
    ]:
        add_number(doc, item)

    add_heading(doc, "2. Literature Review and Research Gap", 1)
    lit_paras = [
        "Lean manufacturing has long been associated with better operational performance, especially in terms of productivity, quality, throughput, and waste reduction. Belekoukias et al. [1] showed that lean tools can improve manufacturing performance, but they also emphasized that their benefits depend strongly on implementation context. In other words, lean methods are not universally equal in impact; the best tool depends on the specific nature of losses in the system.",
        "Quantitative lean analysis therefore becomes important. Oleghe and Salonitis [2] argued that the interaction between human systems, process variability, and lean implementation still requires stronger analytical modeling. Their work highlights the importance of moving beyond descriptive lean case studies toward more systematic and measurable approaches.",
        "Simulation has emerged as one of the most suitable techniques for that purpose. Yu and Chen [3], in their review of simulation use in SMEs, found that factory simulation offers strong potential for improvement planning but remains underused in smaller firms. They identified the need for more practical, accessible, and transferable simulation applications suited to SME environments.",
        "Possik et al. [4] developed a co-simulation approach to evaluate lean techniques in manufacturing systems and showed that different tools produce different levels of effect depending on system design. Frecassetti et al. [5] demonstrated in an Italian SME that simulation can support the introduction of lean practices by helping managers compare alternatives before implementation.",
        "At the same time, lean and digital manufacturing have increasingly converged. Bokhorst et al. [6] showed that smart manufacturing often builds on lean principles rather than replacing them. More recent work on digital twins and automated value stream mapping also indicates that real-time digital data can strengthen lean diagnosis and decision-making [10].",
        "In food and beverage SMEs, lean implementation is also receiving growing attention. Matindana and Shoshiwa [9] found that lean practices are relevant in such environments, but quantitative studies with event-level data remain comparatively limited. This is significant because food and beverage systems often experience short interruptions, changeovers, maintenance losses, and operational instability that require targeted analysis.",
        "Based on the literature, the main research gap can be summarized in three points: many lean-simulation studies rely on proprietary case-company data and are difficult to reproduce; many studies discuss lean conceptually but do not connect event-level downtime records with simulation-based decision analysis; and SME-relevant sectors such as food and beverage still have relatively fewer open, data-driven case studies compared to larger industrial sectors.",
    ]
    for para in lit_paras:
        add_paragraph(doc, para)
    add_table(
        doc,
        "Table 1. Literature Synthesis and Research Gap",
        ["Study", "Main Contribution", "Gap Remaining"],
        [
            ["Yu and Chen [3]", "Review of factory simulation in SMEs and identification of adoption barriers.", "Need for more transferable and practical SME case studies."],
            ["Possik et al. [4]", "Co-simulation framework for evaluating lean technique impact.", "Limited open-data reproducibility."],
            ["Frecassetti et al. [5]", "Simulation-based introduction of lean practices in an SME.", "Industrial data not openly reusable."],
            ["Bokhorst et al. [6]", "Demonstrates the relation between smart manufacturing and lean principles.", "Does not directly evaluate downtime-driven lean strategies."],
            ["Reslan et al. [10]", "Digital twin support for value stream mapping.", "Limited productivity simulation from open event data."],
            ["This study", "Open-data, downtime-driven simulation of lean alternatives.", "Provides a reproducible framework for student and SME use."],
        ],
    )

    add_heading(doc, "3. Problem Definition and Mathematical Formulation", 1)
    for para in [
        "The production line studied in this paper suffers from large variations in operating time and downtime. The central problem is to determine which lean strategy should be given priority in order to improve productivity most effectively.",
        "The key relationships used in this study are listed below in Word-ready form.",
    ]:
        add_paragraph(doc, para)
    for equation in [
        "Operational Availability: A = T_op / T_mon",
        "Downtime: T_dt = T_mon - T_op",
        "Production Rate: R = Q / T_op",
        "Scenario-Based Output: Q' = R x (T_mon - T'_dt)",
        "Downtime Transformation: T'_dt = T_res + Σ(alpha_i x t_i)",
        "Percentage Productivity Improvement: %ΔQ = ((Q' - Q) / Q) x 100",
    ]:
        add_paragraph(doc, equation, bold=True)

    add_heading(doc, "4. Dataset Description and Study Context", 1)
    for para in [
        "The dataset used in this study is the Industrial Production Time-Series Dataset from a Beverage Bottling Line, published on Zenodo on January 4, 2026 [8]. The data correspond to a real industrial bottling-line operation monitored through an IIoT architecture. The associated research article by David et al. [7] confirms the industrial and digital monitoring context of the data source.",
        "The dataset contains four important groups of information: production output aggregated by hour, daily production summaries, hourly operating breakdown records, and event-level downtime logs.",
        "The monitoring period covers July 22, 2022 to February 6, 2023. The dataset includes two product sizes, 3-liter and 5-liter containers, and captures real operating instability over time. Because the data are open and reusable, they are highly suitable for an academic study requiring transparency and reproducibility.",
    ]:
        add_paragraph(doc, para)
    add_table(
        doc,
        "Table 2. Dataset Characteristics",
        ["Attribute", "Value"],
        [
            ["Monitoring period", "July 22, 2022 to February 6, 2023"],
            ["Production summary records", "59"],
            ["Calendar production days", "56"],
            ["Hourly operating observations", "265"],
            ["Downtime events", "1,388"],
            ["Observed product sizes", "3 L and 5 L"],
            ["Total liters produced", "212,358"],
            ["Total units produced", "60,888"],
            ["Total monitored time", "238.65 h"],
            ["Total operating time", "101.81 h"],
        ],
    )

    add_heading(doc, "5. Research Methodology", 1)
    methodology = {
        "Step 1: Data Extraction and Cleaning": "The Excel dataset was first read and organized into production summaries, hourly operation records, and downtime-event data. Dates, durations, and production quantities were standardized. Downtime durations were converted to minutes so that event-level categorization and simulation could be performed consistently.",
        "Step 2: Baseline Performance Analysis": "Baseline performance was measured using liters produced, production units, monitored time, operating time, pause time, daily efficiency, and hourly efficiency. This step was necessary to understand the real operating condition of the system before any lean scenario was introduced.",
        "Step 3: Downtime Categorization": "Downtime events were classified into three categories: micro-stoppages (t <= 2 min), minor stoppages (2 < t <= 10 min), and major stoppages (t > 10 min). This classification distinguishes frequent short interruptions from serious breakdown-type events.",
        "Step 4: Production-Rate Estimation": "For each observed production day, production rate was estimated from liters produced divided by effective operating time. This creates a realistic day-level performance profile that preserves operational variability instead of assuming a constant theoretical output rate.",
        "Step 5: Bootstrap Simulation": "An event-level bootstrap simulation with 5,000 replications was performed. Each replication resampled observed day profiles with replacement. This allows the model to preserve the stochastic behavior of monitored time, productivity, and downtime structure. Event durations were normalized to the recorded daily pause time so that the simulated baseline remained aligned with the observed plant record.",
        "Step 6: Scenario Design": "Three lean scenarios were designed and evaluated: 5S + Visual Management with 15% of micro-stoppages eliminated; TPM with 25% reduction in the duration of major stoppages; and an Integrated Lean Bundle with 20% of micro-stoppages eliminated, 15% reduction in minor stoppages, and 30% reduction in major stoppages. These values are scenario assumptions informed by lean implementation logic and literature.",
    }
    for title, text in methodology.items():
        add_paragraph(doc, title, bold=True)
        add_paragraph(doc, text)

    add_paragraph(doc, "Algorithm and Logic Used in the Model", bold=True)
    add_paragraph(doc, "The main algorithm used in this research is a data-driven bootstrap simulation. First, the real production and downtime records are converted into daily operating profiles. Second, these real daily profiles are resampled repeatedly so that actual plant variability is preserved. Third, lean-improvement rules are applied to the sampled downtime events. Finally, output, pause time, and efficiency are recalculated.")
    add_paragraph(doc, "In simple words, the algorithm repeatedly asks: if we reduce the type of downtime targeted by a lean tool, how much production can be recovered?")
    for item in [
        "Read the production sheets and downtime-event sheet.",
        "Convert downtime durations into minutes and aggregate them by date.",
        "Calculate daily monitored time, operating time, pause time, and production rate.",
        "Create baseline operating profiles for each observed day.",
        "Sample those observed day profiles with replacement across many replications.",
        "Apply lean rules to the sampled downtime events.",
        "Recalculate pause time, output, and efficiency after each scenario.",
        "Compare the simulated results with the baseline system.",
    ]:
        add_number(doc, item)

    add_paragraph(doc, "How Lean Manufacturing Was Applied to the Data", bold=True)
    add_paragraph(doc, "Lean manufacturing was used directly on the observed downtime data instead of being discussed only as a theory. In this study, downtime was treated as measurable waste, and each lean tool was linked to a specific category of loss in the dataset.")
    add_table(
        doc,
        "Table 2A. How Lean Tools Were Mapped to the Data",
        ["Lean Tool", "Data Feature Used", "Logic Used in the Study"],
        [
            ["5S + Visual Management", "Micro-stoppages", "Used to represent the reduction of small avoidable interruptions caused by poor workplace organization and visual control."],
            ["TPM", "Major stoppages", "Used to represent reduced machine breakdown duration and faster recovery through better maintenance."],
            ["Integrated Lean Bundle", "Micro, minor, and major stoppages", "Used to represent a combined improvement in workplace discipline and machine reliability."],
        ],
    )

    add_heading(doc, "6. Results and Visual Analysis", 1)
    add_paragraph(doc, "The observed production system produced 212,358 liters from 60,888 units during 238.65 monitored hours. Effective operating time was only 101.81 hours, while pause time reached 137.56 hours. This corresponds to an overall observed availability of 42.66%, indicating that downtime is the main performance constraint.")
    add_paragraph(doc, "At the product level, the average production rate was 1,963.71 L/h for the 3-liter product and 2,918.68 L/h for the 5-liter product. The 5-liter product therefore showed a higher effective throughput during operating time.")
    add_table(
        doc,
        "Table 3. Baseline Performance Indicators",
        ["Metric", "Value"],
        [
            ["Production days", "56"],
            ["Production summary records", "59"],
            ["Hourly observations", "265"],
            ["Downtime events", "1,388"],
            ["Total production", "212,358 L"],
            ["Total units produced", "60,888"],
            ["Monitored time", "238.65 h"],
            ["Operating time", "101.81 h"],
            ["Pause time", "137.56 h"],
            ["Overall availability", "42.66%"],
        ],
    )

    images = [
        ("monthly_efficiency.png", "Figure 1. Monthly average operational efficiency from the real bottling-line dataset.", 6.2),
        ("downtime_distribution.png", "Figure 2. Distribution of downtime-event duration, truncated at 60 minutes for readability.", 6.2),
    ]
    for filename, caption, width in images:
        add_picture(doc, ROOT / "outputs" / filename, caption, width)
        doc.add_paragraph()

    add_table(
        doc,
        "Table 4. Downtime Category Structure",
        ["Category", "Event Count", "Share of Downtime Minutes"],
        [
            ["Micro (<= 2 min)", "851", "13.59%"],
            ["Minor (2-10 min)", "443", "20.15%"],
            ["Major (> 10 min)", "94", "66.26%"],
        ],
    )

    add_picture(doc, ROOT / "outputs" / "downtime_category_pareto.png", "Figure 3. Downtime Pareto by event category. Major stoppages are rare but dominant in lost time.", 6.2)
    doc.add_paragraph()
    add_paragraph(doc, "The downtime Pareto view provides one of the most important insights of the paper. Micro-stoppages account for 61.31% of all events but only 13.59% of total downtime minutes, whereas major stoppages represent only 6.77% of events but consume 66.26% of lost time. This means that event frequency alone does not identify the highest productivity loss. Lean prioritization must be based on lost-time severity.")

    add_picture(doc, ROOT / "outputs" / "pause_vs_efficiency.png", "Figure 4. Daily pause time versus operational efficiency.", 6.2)
    doc.add_paragraph()
    add_paragraph(doc, "The scatter plot shows a clear downward trend. This provides further evidence that reducing downtime is the most direct productivity-improvement lever in the observed system.")

    add_picture(doc, ROOT / "outputs" / "product_rate_comparison.png", "Figure 5. Mean and median production rate by product size.", 5.8)
    doc.add_paragraph()
    add_paragraph(doc, "The 5-liter product has both higher mean and higher median production rate than the 3-liter product. This suggests that product mix can also influence line performance and may become an additional decision variable in future scheduling or line-balancing studies.")

    add_paragraph(doc, "The simulation baseline produced 211,924.69 liters on average across the resampled production horizon. This closely matches the observed output of the real system, indicating that the simulation reasonably preserves actual operating behavior.")
    add_table(
        doc,
        "Table 5. Simulation Results for Lean Scenarios",
        ["Scenario", "Mean Output (L)", "Gain (%)", "Efficiency (%)", "Eff. Gain (pp)"],
        [
            ["Baseline", "211,924.69", "0.00", "42.37", "0.00"],
            ["5S + Visual Management", "216,866.69", "2.33", "43.25", "0.88"],
            ["TPM", "255,604.10", "20.61", "50.81", "8.44"],
            ["Integrated Lean Bundle", "280,142.22", "32.19", "55.33", "12.96"],
        ],
    )
    add_table(
        doc,
        "Table 6. Simulation Uncertainty Range",
        ["Scenario", "P05 Output (L)", "P95 Output (L)"],
        [
            ["Baseline", "188,572.81", "236,799.79"],
            ["5S + Visual Management", "192,906.87", "242,100.46"],
            ["TPM", "226,295.58", "287,389.24"],
            ["Integrated Lean Bundle", "247,828.10", "314,804.33"],
        ],
    )
    add_picture(doc, ROOT / "outputs" / "scenario_output.png", "Figure 6. Mean simulated production output across the baseline and lean scenarios.", 6.2)
    doc.add_paragraph()
    add_paragraph(doc, "The simulation results show that 5S and visual management provide only a modest output gain of 2.33%. This is because the scenario mainly targets micro-stoppages, which are frequent but contribute relatively little to total lost time. By contrast, TPM produces a much larger gain of 20.61% because it directly targets major stoppages, which dominate downtime minutes. The integrated lean bundle produces the highest gain, improving output by 32.19% and raising simulated availability from 42.37% to 55.33%.")

    add_heading(doc, "7. Discussion", 1)
    for para in [
        "The results strongly suggest that lean implementation should be based on the structure of operational loss rather than on the visibility or frequency of events alone. In many production systems, micro-stoppages attract attention because they happen often and disrupt workflow visibly. However, as demonstrated in this study, those frequent interruptions may not be the dominant cause of lost production.",
        "For the bottling line studied here, major stoppages are the real productivity driver. Although they are relatively rare, they consume most of the downtime minutes. This explains why TPM, which focuses on equipment reliability and breakdown reduction, produces far more improvement than 5S alone. The result is fully consistent with the broader lean-simulation literature, which emphasizes context-specific tool selection [4], [5].",
        "Another important implication is the role of digital manufacturing data in strengthening lean analysis. Instead of relying only on static value stream mapping or qualitative observation, this study uses event-level downtime records and simulation. This is aligned with the direction suggested by smart manufacturing and digital twin research [6], [10]. The practical value is clear: real data allow the analyst to identify high-leverage interventions rather than applying lean tools generically.",
        "For final-year engineering research, this approach is especially useful because it combines core industrial engineering ideas: data analysis, productivity evaluation, lean manufacturing, operations research thinking, and simulation-based decision support. It also produces a study that is more rigorous than a purely descriptive case report.",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "8. Implementation Recommendations", 1)
    add_paragraph(doc, "Based on the results, the recommended order of lean implementation is:")
    for item in [
        "First priority: TPM. The plant should focus on reducing long-duration stoppages through maintenance planning, preventive checks, fast fault recovery, and equipment reliability tracking.",
        "Second priority: 5S and visual management. These tools should be used to reduce small disruptions, improve operator control, and standardize the workplace.",
        "Third priority: integrated lean bundle. Once maintenance discipline is improved, the plant can combine TPM with micro-stop reduction and standard work for broader gains.",
    ]:
        add_number(doc, item)
    add_table(
        doc,
        "Table 7. Recommended Lean Implementation Roadmap",
        ["Priority", "Lean Focus", "Expected Impact"],
        [
            ["1", "TPM and maintenance discipline", "Highest reduction in lost-time severity by targeting major stoppages."],
            ["2", "5S and visual management", "Better workplace organization and reduction in frequent micro-disruptions."],
            ["3", "Integrated lean bundle", "Combined gains from reliability, stabilization, and standardized work."],
        ],
    )

    add_heading(doc, "9. Conclusion", 1)
    for para in [
        "This paper presented a self-contained, one-column research study on simulation-based lean manufacturing using a real industrial beverage bottling dataset. The study showed that the system suffers from low availability, high downtime, and substantial variability over time. Although micro-stoppages are the most frequent events, major stoppages dominate lost production time.",
        "The simulation results demonstrated that 5S and visual management improve output by 2.33%, TPM improves output by 20.61%, and an integrated lean bundle improves output by 32.19% while raising simulated availability from 42.37% to 55.33%.",
        "The main managerial conclusion is that maintenance-led reduction of major stoppages should be prioritized first. The main academic contribution is the development of a reproducible, open-data framework for simulation-based lean research suitable for student projects, SME analysis, and future applied manufacturing studies.",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "10. Limitations and Future Scope", 1)
    for para in [
        "This study has three major limitations. First, the lean scenarios are simulated assumptions rather than observed post-implementation outcomes. Second, the dataset does not contain detailed cause-coded downtime labels, so the mapping from downtime category to lean tool remains inferential. Third, the analysis is based on one industrial case, so the results should not be generalized mechanically to all production systems.",
        "Future work can improve the present study by collecting root-cause-coded downtime data, operator-level observations, quality-loss information, and real post-intervention performance records. Future studies may also extend the simulation to include changeover analysis, workforce allocation, line balancing, and cost-benefit evaluation.",
    ]:
        add_paragraph(doc, para)

    add_heading(doc, "Acknowledgement", 1)
    add_paragraph(doc, "The authors sincerely acknowledge the guidance and support of Dr. MOHD SHUAIB from the Mechanical Department. His academic guidance was valuable in shaping the direction and structure of this work.")

    add_heading(doc, "References", 1)
    references = [
        "[1] I. Belekoukias, J. A. Garza-Reyes, and V. Kumar, “The impact of lean methods and tools on the operational performance of manufacturing organisations,” International Journal of Production Research, vol. 52, no. 18, pp. 5346-5366, 2014. doi: 10.1080/00207543.2014.903348.",
        "[2] O. Oleghe and K. Salonitis, “Hybrid simulation modelling of the human-production process interface in lean manufacturing systems,” International Journal of Lean Six Sigma, vol. 10, no. 2, pp. 665-690, 2019. doi: 10.1108/IJLSS-01-2018-0004.",
        "[3] F. Yu and Z. Chen, “Tools, application areas and challenges of factory simulation in Small and Medium-Sized Enterprises - A Review,” Procedia CIRP, vol. 104, pp. 399-404, 2021. doi: 10.1016/j.procir.2021.11.067.",
        "[4] J. Possik, A. Zouggar-Amrani, B. Vallespir, and G. Zacharewicz, “Lean techniques impact evaluation methodology based on a co-simulation framework for manufacturing systems,” International Journal of Computer Integrated Manufacturing, vol. 35, no. 1, pp. 91-111, 2022. doi: 10.1080/0951192X.2021.1972468.",
        "[5] S. Frecassetti, B. Kassem, K. Kundu, M. Ferrazzi, and A. Portioli-Staudacher, “Introducing Lean practices through simulation: A case study in an Italian SME,” Quality Management Journal, vol. 30, no. 2, pp. 90-104, 2023. doi: 10.1080/10686967.2023.2171326.",
        "[6] J. A. C. Bokhorst, W. Knol, J. Slomp, and T. Bortolotti, “Assessing to what extent smart manufacturing builds on lean principles,” International Journal of Production Economics, vol. 253, art. no. 108599, 2022. doi: 10.1016/j.ijpe.2022.108599.",
        "[7] G. A. David, P. M. C. Monson, C. Soares Junior, P. O. Conceicao Junior, P. R. Aguiar, and A. Simeone, “IoT-Driven Deep Learning for Enhanced Industrial Production Forecasting,” IEEE Internet of Things Journal, vol. 11, no. 23, pp. 38486-38495, 2024. doi: 10.1109/JIOT.2024.3447579.",
        "[8] G. A. David, P. M. C. Monson, and C. Soares Junior, Industrial Production Time-Series Dataset from a Beverage Bottling Line. Zenodo, Jan. 4, 2026. doi: 10.5281/zenodo.18146866.",
        "[9] J. M. Matindana and M. J. Shoshiwa, “Lean manufacturing implementation in food and beverage SMEs in Tanzania: using structural equation modelling (SEM),” Management System Engineering, vol. 4, no. 1, pp. 1-14, 2025. doi: 10.1007/s44176-025-00036-3.",
        "[10] M. Reslan, M. J. Triebe, R. Venketesh, and A. J. Hartwell, “Automation of Value Stream Mapping: A Case Study on Enhancing Lean Manufacturing Tools Through Digital Twins,” Procedia CIRP, vol. 134, pp. 455-460, 2025. doi: 10.1016/j.procir.2025.02.159.",
        "[11] C. Unal and S. Bilget, “Examination of lean manufacturing systems by simulation technique in apparel industry,” The Journal of The Textile Institute, vol. 112, no. 3, pp. 377-387, 2021. doi: 10.1080/00405000.2020.1756104.",
    ]
    for ref in references:
        add_paragraph(doc, ref)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(f"Created: {OUTPUT}")
